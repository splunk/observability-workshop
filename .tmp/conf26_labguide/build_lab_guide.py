from copy import deepcopy
from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


REFERENCE = Path(
    "/Users/kylwang/Downloads/"
    "conf26 TEMPLATE - [Workshop Name] - Lab Guide - [Date] (1).docx"
)
OUTPUT = Path(
    "/Users/kylwang/Documents/Workshop v2/"
    "conf26 OBS1184 - Advanced OpenTelemetry Collector - Lab Guide.docx"
)
WORKING = Path(
    "/Users/kylwang/Documents/Workshop v2/.tmp/conf26_labguide/work/"
    "lab-guide-working.docx"
)


def full_text(paragraph):
    return "".join(node.text or "" for node in paragraph._p.xpath(".//w:t"))


def paragraphs_by_exact_text(doc, text):
    return [p for p in doc.paragraphs if full_text(p) == text]


def clear_paragraph_content(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_text_run(paragraph, text, bold=False, italic=False, code=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Consolas"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), "Consolas")
    return run


def add_inline_markup(paragraph, text):
    # Minimal parser for **bold** and `code` used by the authored content.
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            add_text_run(paragraph, part[2:-2], bold=True)
        elif part.startswith("`") and part.endswith("`"):
            add_text_run(paragraph, part[1:-1], code=True)
        else:
            add_text_run(paragraph, part)


def move_before(paragraph, marker_el):
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)
    marker_el.addprevious(paragraph._p)
    return paragraph


def add_body(doc, marker_el, text, indent=0.0, after=5, keep_next=False):
    paragraph = doc.add_paragraph(style="Normal")
    clear_paragraph_content(paragraph)
    add_inline_markup(paragraph, text)
    paragraph.paragraph_format.left_indent = Inches(indent)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = keep_next
    return move_before(paragraph, marker_el)


def add_lead(doc, marker_el, text):
    paragraph = doc.add_paragraph(style="Normal")
    clear_paragraph_content(paragraph)
    run = add_text_run(paragraph, text, bold=True)
    run.font.size = Pt(12)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    return move_before(paragraph, marker_el)


def clone_clean_paragraph(doc, source_p, marker_el):
    clone = deepcopy(source_p._p)
    for child in list(clone):
        if child.tag != qn("w:pPr"):
            clone.remove(child)
    marker_el.addprevious(clone)
    from docx.text.paragraph import Paragraph

    return Paragraph(clone, doc._body)


def add_numbered_step(doc, marker_el, source_p, text):
    paragraph = clone_clean_paragraph(doc, source_p, marker_el)
    add_inline_markup(paragraph, text)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet(doc, marker_el, source_p, text, indent=0.55):
    paragraph = clone_clean_paragraph(doc, source_p, marker_el)
    add_inline_markup(paragraph, text)
    paragraph.paragraph_format.left_indent = Inches(indent)
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def append_code_text(run_el, text):
    for index, line in enumerate(text.split("\n")):
        if index:
            run_el.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        run_el.append(t)


def add_code(doc, marker_el, source_p, text):
    paragraph = clone_clean_paragraph(doc, source_p, marker_el)
    run_el = OxmlElement("w:r")
    source_run = source_p._p.find(qn("w:r"))
    if source_run is not None and source_run.find(qn("w:rPr")) is not None:
        run_el.append(deepcopy(source_run.find(qn("w:rPr"))))
    append_code_text(run_el, text)
    paragraph._p.append(run_el)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_note(doc, marker_el, source_p, text, label="Note: "):
    paragraph = clone_clean_paragraph(doc, source_p, marker_el)
    add_text_run(paragraph, label, bold=True)
    add_inline_markup(paragraph, text)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_hyperlink(paragraph, label, url):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_link_line(doc, marker_el, label, url, lead=None):
    paragraph = doc.add_paragraph(style="Normal")
    clear_paragraph_content(paragraph)
    if lead:
        add_text_run(paragraph, lead, bold=True)
    add_hyperlink(paragraph, label, url)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    return move_before(paragraph, marker_el)


def remove_between(start_el, end_el):
    current = start_el.getnext()
    while current is not None and current is not end_el:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt


def replace_first_text_node(paragraph, old, new):
    for node in paragraph._p.xpath(".//w:t"):
        if node.text == old:
            node.text = new
            return True
    return False


def replace_toc_visible_text(doc, old, new):
    for paragraph in doc.paragraphs:
        if paragraph.style.name.lower().startswith("toc"):
            for node in paragraph._p.xpath(".//w:t"):
                if node.text == old:
                    node.text = new
                    return True
    return False


def add_page_break_before(paragraph):
    paragraph.paragraph_format.page_break_before = True


WORKING.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(REFERENCE, WORKING)
doc = Document(WORKING)

# Capture template components before edits.
all_paragraphs = doc.paragraphs
cover_session = all_paragraphs[0]
cover_title = all_paragraphs[1]
cover_instruction = all_paragraphs[3]
bullet_template = all_paragraphs[26]
exercise_headings = [p for p in all_paragraphs if p.style.name == "Heading 1"][-3:]
description_paragraphs = [all_paragraphs[71], all_paragraphs[81], all_paragraphs[92]]
steps_headings = [all_paragraphs[72], all_paragraphs[82], all_paragraphs[93]]
step_templates = [all_paragraphs[73], all_paragraphs[83], all_paragraphs[94]]
code_template = all_paragraphs[74]
note_template = all_paragraphs[77]

# Cover slots.
clear_paragraph_content(cover_session)
session_run = add_text_run(cover_session, "OBS1184", bold=True)
session_run.font.size = Pt(24)

clear_paragraph_content(cover_title)
title_run = add_text_run(cover_title, "Advanced OpenTelemetry Collector Lab Guide")
title_run.font.size = Pt(16)

clear_paragraph_content(cover_instruction)
meta_run = add_text_run(
    cover_instruction,
    "55-minute hands-on workshop | Splunk Distribution of OpenTelemetry Collector 0.157.0",
    bold=True,
)
meta_run.font.color.rgb = RGBColor(38, 38, 38)
meta_run.font.size = Pt(10.5)

# Update the three existing TOC slots without disturbing PAGEREF fields.
toc_titles = [
    "Exercise 1 - Set Up and Validate the Agent",
    "Exercise 2 - Build the Configuration in Config Builder",
    "Exercise 3 - Deploy and Prove the Changes",
]
for old, new in zip(
    [
        "Exercise 1 – [Title]",
        "Exercise 2 – [Title]",
        "Exercise 3 – [Title]",
    ],
    toc_titles,
):
    assert replace_toc_visible_text(doc, old, new), old

# Update heading text inside the existing bookmarks.
heading_suffixes = [
    " 1 - Set Up and Validate the Agent",
    " 2 - Build the Configuration in Config Builder",
    " 3 - Deploy and Prove the Changes",
]
for exercise_index, (paragraph, old_suffix, new_suffix) in enumerate(zip(
    exercise_headings,
    [" 1 – [Title]", " 2 – [Title]", " 3 – [Title]"],
    heading_suffixes,
)):
    assert replace_first_text_node(paragraph, old_suffix, new_suffix), full_text(paragraph)
    if exercise_index != 1:
        add_page_break_before(paragraph)

# Fill descriptions.
descriptions = [
    (
        "Prepare either a provided Linux workshop instance or an Apple Silicon Mac, "
        "run one Splunk OpenTelemetry Collector in Agent mode, generate baseline "
        "traces and host metrics, and import the starter configuration into OTel "
        "Collector Config Builder. The Collector is pinned to version 0.157.0."
    ),
    (
        "Use the guided Config Builder interface to understand the imported Agent "
        "configuration and add three processing scenarios: drop noisy health spans, "
        "protect synthetic sensitive attributes, and transform structured logs. Keep "
        "all existing receivers, processors, exporters, extensions, and pipelines."
    ),
    (
        "Download the completed YAML once, replace the running Agent configuration, "
        "and compare the original load-generator telemetry with the processed output. "
        "Validate traces and host metrics in Splunk Observability Cloud when cloud "
        "export is enabled. Logs and profiling remain take-home extensions."
    ),
]
for paragraph, text in zip(description_paragraphs, descriptions):
    clear_paragraph_content(paragraph)
    add_inline_markup(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(8)

# Remove all placeholder content after each Steps heading while preserving the next
# exercise heading and the final section properties.
body = doc._body._element
markers = [exercise_headings[1]._p, exercise_headings[2]._p, body.sectPr]
for steps_heading, marker in zip(steps_headings, markers):
    remove_between(steps_heading._p, marker)


def build_exercise_one(marker):
    step = step_templates[0]

    add_numbered_step(
        doc,
        marker,
        step,
        "Pair up if helpful, register for a free Splunk Observability Cloud organization, "
        "and open the online workshop page supplied for session OBS1184 in Google Chrome.",
    )
    add_link_line(
        doc,
        marker,
        "Register for Splunk Observability Cloud Free",
        "https://www.splunk.com/en_us/download/observability-cloud-free-edition.html",
        lead="Registration: ",
    )
    add_body(
        doc,
        marker,
        "Have your Observability Cloud **realm** and an **ingest access token** ready. "
        "The setup script stores them in `workshop-env.sh`; never share or commit that file.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Choose one supported execution path and confirm that `jq` is available with "
        "`jq --version`.",
    )
    add_bullet(
        doc,
        marker,
        bullet_template,
        "**Provided Linux instance:** use the Splunk Show instructions earlier in this "
        "guide, then connect with the SSH host, user, and port supplied for the session. "
        "Outbound access to port 2222 is required.",
    )
    add_code(doc, marker, code_template, "ssh -p 2222 <workshop-user>@<workshop-host>")
    add_bullet(
        doc,
        marker,
        bullet_template,
        "**Apple Silicon Mac:** use Terminal locally. Apple Silicon supports the same "
        "local and cloud-connected exercises.",
    )
    add_bullet(
        doc,
        marker,
        bullet_template,
        "**Windows or Intel Mac:** use the provided Linux workshop instance; the local "
        "setup does not support these platforms.",
    )

    add_numbered_step(doc, marker, step, "Create the workshop directory.")
    add_code(
        doc,
        marker,
        code_template,
        "mkdir advanced-otel-workshop && \\\ncd advanced-otel-workshop",
    )
    add_body(
        doc,
        marker,
        "This guide uses `[WORKSHOP]` for the full path to this directory.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Define the workshop download locations, then download the pinned Collector, "
        "the original workshop `loadgen`, and the setup script.",
    )
    add_code(
        doc,
        marker,
        code_template,
        "BRANCH=codex/advanced-collector-conf2026\n"
        "REPO=https://github.com/chentaow-splunk/observability-workshop/raw/refs/heads/$BRANCH\n"
        "LAB=content/en/ninja-workshops/foundations/2-opentelemetry-collector-workshops/conf2026-obs1184",
    )
    add_lead(doc, marker, "Provided Linux instance")
    add_code(
        doc,
        marker,
        code_template,
        "curl -L https://github.com/signalfx/splunk-otel-collector/releases/download/v0.157.0/otelcol_linux_amd64 -o otelcol\n"
        "curl -L \"$REPO/workshop/ninja/advanced-otel/conf2026-obs1184/loadgen/build/loadgen-linux-amd64\" -o loadgen\n"
        "curl -L \"$REPO/$LAB/setup-workshop-conf2026.sh\" -o setup-workshop.sh\n"
        "chmod +x setup-workshop.sh",
    )
    add_lead(doc, marker, "Apple Silicon Mac")
    add_code(
        doc,
        marker,
        code_template,
        "curl -L https://github.com/signalfx/splunk-otel-collector/releases/download/v0.157.0/otelcol_darwin_arm64 -o otelcol\n"
        "curl -L \"$REPO/workshop/ninja/advanced-otel/conf2026-obs1184/loadgen/build/loadgen-darwin-arm64\" -o loadgen\n"
        "curl -L \"$REPO/$LAB/setup-workshop-conf2026.sh\" -o setup-workshop.sh\n"
        "chmod +x setup-workshop.sh",
    )

    add_numbered_step(doc, marker, step, "Run the setup script.")
    add_code(doc, marker, code_template, "./setup-workshop.sh")
    add_body(
        doc,
        marker,
        "Press Enter at the first prompt for the default cloud-connected mode. Enter "
        "your realm and ingest token, and accept the proposed API URL unless your "
        "organization specifies another endpoint. Press Enter to skip both optional "
        "HEC prompts; HEC is not required for the live workshop.",
        indent=0.25,
    )
    add_body(
        doc,
        marker,
        "Choose local-only mode only when you intentionally do not want to send traces "
        "and metrics to Splunk Observability Cloud.",
        indent=0.25,
    )

    add_numbered_step(doc, marker, step, "Verify that setup created the Agent files.")
    add_code(
        doc,
        marker,
        code_template,
        "test -f 1-agent/agent_config.yaml && \\\ntest -f workshop-env.sh && \\\necho \"Workshop setup is complete.\"",
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Open three terminals named **Agent Console**, **Loadgen**, and **Tests**. In "
        "Agent Console, start the Collector and leave it running.",
    )
    add_code(
        doc,
        marker,
        code_template,
        "cd [WORKSHOP]/1-agent\n"
        "source ../workshop-env.sh\n"
        "../otelcol --config=agent_config.yaml",
    )
    add_body(
        doc,
        marker,
        "Confirm that startup reaches `Everything is ready. Begin running and processing data.`",
        indent=0.25,
    )

    add_numbered_step(doc, marker, step, "In Tests, verify the Agent health endpoint.")
    add_code(doc, marker, code_template, "curl -fsS http://127.0.0.1:13133/")

    add_numbered_step(doc, marker, step, "Verify baseline host metrics and traces.")
    add_body(
        doc,
        marker,
        "After at least 10 seconds, the Agent Console should show `system.*` host "
        "metrics. Linux commonly includes CPU, memory, load, and network metrics; "
        "Apple Silicon can expose a different OS-supported subset.",
        indent=0.25,
    )
    add_code(doc, marker, code_template, "cd [WORKSHOP]/1-agent\n../loadgen -count 5")
    add_body(
        doc,
        marker,
        "Confirm five `/movie-validator` spans in the Agent Console. The baseline span "
        "contains synthetic `user.*` attributes that you will protect later.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "If cloud export is enabled, confirm the baseline in Splunk Observability Cloud.",
    )
    add_body(
        doc,
        marker,
        "Open **APM > Trace Analyzer**, select **All traces**, and find service "
        "`cinema-service` with operation `/movie-validator`. Then open "
        "**Infrastructure > Hosts** and locate the detected host. Cloud-provider and "
        "Kubernetes hosts can appear in their platform-specific navigator.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Import the starter `agent_config.yaml` into **Data Management > OTel Collector "
        "Config Builder**. Review Component Inventory, Pipelines, and Collector YAML.",
    )
    add_body(
        doc,
        marker,
        "The imported file combines Splunk Distribution defaults with workshop-only "
        "debug and file exporters. It uses one Agent and no Gateway. If the workshop "
        "runs remotely, download the same starter YAML from the online workshop page "
        "to the computer running your browser before importing it.",
        indent=0.25,
    )

    add_note(
        doc,
        marker,
        note_template,
        "Keep `workshop-env.sh` and all tokens outside Config Builder. Import only "
        "`agent_config.yaml`, which contains environment-variable references rather "
        "than credential values.",
    )


def build_exercise_two(marker):
    step = step_templates[1]

    add_numbered_step(
        doc,
        marker,
        step,
        "In **Component Inventory**, add processor `filter` with component name "
        "`health`, producing component ID `filter/health`.",
    )
    add_body(
        doc,
        marker,
        "Set top-level `error_mode` to `ignore`. Under `trace_conditions`, add one "
        "condition group and enter:",
        indent=0.25,
    )
    add_code(doc, marker, code_template, 'span.name == "/_healthz"')
    add_body(
        doc,
        marker,
        "Leave the condition group's context and error mode empty. Preview the YAML "
        "and add the component. Use `trace_conditions`; the legacy `traces.span` form "
        "is deprecated.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Edit the `traces` pipeline and place `filter/health` immediately after "
        "`memory_limiter`. Keep every existing receiver, processor, and exporter.",
    )
    add_code(
        doc,
        marker,
        code_template,
        "processors:\n"
        "  - memory_limiter\n"
        "  - filter/health\n"
        "  - resource/add_mode\n"
        "  - batch\n"
        "  - resource_detection",
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Add processor `attributes`. In Options, create these actions in order.",
    )
    add_bullet(doc, marker, bullet_template, "Update `user.phone_number` to `UNKNOWN NUMBER`.")
    add_bullet(doc, marker, bullet_template, "Hash `user.email`.")
    add_bullet(doc, marker, bullet_template, "Delete `user.password`.")
    add_code(
        doc,
        marker,
        code_template,
        "attributes:\n"
        "  actions:\n"
        "    - key: user.phone_number\n"
        "      action: update\n"
        "      value: UNKNOWN NUMBER\n"
        "    - key: user.email\n"
        "      action: hash\n"
        "    - key: user.password\n"
        "      action: delete",
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Add processor `redaction`. Set `allow_all_keys` to **True**, add the Visa and "
        "Mastercard patterns below to `blocked_values`, and set `summary` to `debug`.",
    )
    add_code(
        doc,
        marker,
        code_template,
        r"\b4[0-9]{3}[\s-]?[0-9]{4}[\s-]?[0-9]{4}[\s-]?[0-9]{4}\b"
        "\n"
        r"\b5[1-5][0-9]{2}[\s-]?[0-9]{4}[\s-]?[0-9]{4}[\s-]?[0-9]{4}\b",
    )
    add_body(
        doc,
        marker,
        "The Amex value is intentionally unmatched so validation demonstrates an "
        "incomplete policy. Leave the database sanitizer and other optional fields unset.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Add `attributes` and `redaction` to the `traces` pipeline after "
        "`filter/health` and before `resource/add_mode`.",
    )
    add_code(
        doc,
        marker,
        code_template,
        "traces:\n"
        "  receivers: [jaeger, otlp, zipkin]\n"
        "  processors: [memory_limiter, filter/health, attributes, redaction, resource/add_mode, batch, resource_detection]\n"
        "  exporters: [otlp_http, debug, file/traces]",
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Add processor `transform`. Set `error_mode` to `ignore`. Under "
        "`log_statements`, add a `resource` context containing:",
    )
    add_code(
        doc,
        marker,
        code_template,
        'keep_keys(resource.attributes, ["com.splunk.sourcetype", "host.name", "otelcol.service.mode"])',
    )
    add_body(
        doc,
        marker,
        "Add a second `log` context with no statements, preview the component, and "
        "click **Add**. The resource statement removes metadata that is not required "
        "after collection.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Open **Collector YAML > Edit YAML**. Under the existing `- context: log` "
        "entry, paste this indented `statements` block and save only after the editor "
        "shows **VALID**.",
    )
    add_code(
        doc,
        marker,
        code_template,
        "        statements:\n"
        "          - set(log.cache, ParseJSON(log.body)) where IsMatch(log.body, \"^\\\\{\")\n"
        "          - flatten(log.cache, \"\")\n"
        "          - merge_maps(log.attributes, log.cache, \"upsert\")\n"
        "          - set(log.severity_text, log.attributes[\"level\"])\n"
        "          - set(log.severity_number, 1) where log.severity_text == \"TRACE\"\n"
        "          - set(log.severity_number, 5) where log.severity_text == \"DEBUG\"\n"
        "          - set(log.severity_number, 9) where log.severity_text == \"INFO\"\n"
        "          - set(log.severity_number, 13) where log.severity_text == \"WARN\"\n"
        "          - set(log.severity_number, 17) where log.severity_text == \"ERROR\"\n"
        "          - set(log.severity_number, 21) where log.severity_text == \"FATAL\"",
    )
    add_body(
        doc,
        marker,
        "`ParseJSON` creates a temporary map, `flatten` normalizes nested fields, and "
        "`merge_maps(..., \"upsert\")` promotes the parsed fields to log attributes. "
        "The remaining statements map application levels to OpenTelemetry severity.",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "Edit the `logs` pipeline and add `transform` after `resource_detection`. "
        "Keep all existing receivers and exporters.",
    )
    add_code(
        doc,
        marker,
        code_template,
        "logs:\n"
        "  receivers: [fluent_forward, otlp, file_log/quotes]\n"
        "  processors: [memory_limiter, resource/add_mode, batch, resource_detection, transform]\n"
        "  exporters: [splunk_hec, splunk_hec/profiling, debug, file/logs]",
    )

    add_numbered_step(doc, marker, step, "Review the final generated YAML before download.")
    add_bullet(doc, marker, bullet_template, "`filter/health`, `attributes`, and `redaction` appear exactly once in `traces.processors`.")
    add_bullet(doc, marker, bullet_template, "`transform` appears exactly once in `logs.processors`, after `resource_detection`.")
    add_bullet(doc, marker, bullet_template, "The six imported pipelines and all existing components remain present.")
    add_bullet(doc, marker, bullet_template, "Cloud and local exporters remain connected to the traces and metrics pipelines.")

    add_note(
        doc,
        marker,
        note_template,
        "Processor order matters. Filtering first prevents dropped spans from reaching "
        "later processors. Resource detection must run before the log resource allowlist "
        "so `host.name` is available to keep. Promoting every JSON field can create high "
        "cardinality in production; use an explicit field allowlist for real workloads.",
    )


def build_exercise_three(marker):
    step = step_templates[2]

    add_numbered_step(doc, marker, step, "Download the completed Collector YAML as `agent_config.yaml`.")
    add_body(
        doc,
        marker,
        "Confirm that the file contains environment-variable references such as "
        "`${SPLUNK_ACCESS_TOKEN}`, not real credentials.",
        indent=0.25,
    )

    add_numbered_step(doc, marker, step, "Stop the Agent with Ctrl-C, then preserve the starter configuration.")
    add_code(
        doc,
        marker,
        code_template,
        "cd [WORKSHOP]/1-agent\n"
        "if [ ! -f agent_config.start.yaml ]; then\n"
        "  cp agent_config.yaml agent_config.start.yaml\n"
        "fi",
    )

    add_numbered_step(doc, marker, step, "Replace the Agent configuration using the path that matches your environment.")
    add_lead(doc, marker, "Apple Silicon or same computer")
    add_code(
        doc,
        marker,
        code_template,
        "cp ~/Downloads/agent_config.yaml [WORKSHOP]/1-agent/agent_config.yaml",
    )
    add_lead(doc, marker, "Remote workshop instance - run from your local computer")
    add_code(
        doc,
        marker,
        code_template,
        "scp -P 2222 ~/Downloads/agent_config.yaml \\\n  <workshop-user>@<workshop-host>:~/advanced-otel-workshop/1-agent/agent_config.yaml",
    )

    add_numbered_step(doc, marker, step, "Move the earlier quote log and restart the updated Agent.")
    add_code(
        doc,
        marker,
        code_template,
        "cd [WORKSHOP]/1-agent\n"
        "if [ -f quotes.log ]; then\n"
        "  mv quotes.log quotes.log.before-config-builder\n"
        "fi\n"
        "source ../workshop-env.sh\n"
        "../otelcol --config=agent_config.yaml",
    )
    add_body(
        doc,
        marker,
        "Confirm that startup again reaches `Everything is ready. Begin running and processing data.`",
        indent=0.25,
    )

    add_numbered_step(
        doc,
        marker,
        step,
        "In Loadgen, reuse the original workshop generator to create five application "
        "spans and five health-check spans. Copy one base trace ID.",
    )
    add_code(doc, marker, code_template, "cd [WORKSHOP]/1-agent\n../loadgen -health -count 5")

    add_numbered_step(doc, marker, step, "Compare the original and processed span counts locally.")
    add_bullet(doc, marker, bullet_template, "Original from `loadgen`: five `/movie-validator` and five `/_healthz` spans.")
    add_bullet(doc, marker, bullet_template, "Expected after the Collector: five `/movie-validator` spans and zero `/_healthz` spans.")
    add_code(
        doc,
        marker,
        code_template,
        "jq -s -r '\n"
        "  [.[].resourceSpans[].scopeSpans[].spans[]]\n"
        "  | group_by(.name)[]\n"
        "  | \"\\(length) \\(.[0].name)\"\n"
        "' agent-traces.out",
    )
    add_body(doc, marker, "Expected output: `5 /movie-validator`.", indent=0.25)

    add_numbered_step(doc, marker, step, "Compare the original and protected span attributes.")
    add_bullet(doc, marker, bullet_template, "Phone: `+1555-867-5309` -> `UNKNOWN NUMBER`.")
    add_bullet(doc, marker, bullet_template, "Email: `george@deathstar.email` -> deterministic SHA-256 hash.")
    add_bullet(doc, marker, bullet_template, "Password: `LOTR>StarWars1-2-3` -> deleted.")
    add_bullet(doc, marker, bullet_template, "Visa and Mastercard -> `****`; Amex remains visible by design.")
    add_code(
        doc,
        marker,
        code_template,
        "jq -s '[\n"
        "  .[].resourceSpans[].scopeSpans[].spans[].attributes[]\n"
        "  | select(.key == \"user.password\")\n"
        "] | length' agent-traces.out",
    )
    add_body(doc, marker, "Expected password count: `0`.", indent=0.25)

    add_numbered_step(
        doc,
        marker,
        step,
        "When cloud export is enabled, prove the trace changes in Splunk APM.",
    )
    add_body(
        doc,
        marker,
        "Open **APM > Trace Analyzer**, select **All traces**, and choose a time range "
        "beginning just before the test. Search for the copied trace ID. The waterfall "
        "must contain `/movie-validator` and must not contain `/_healthz`. Select the "
        "application span and inspect its properties to confirm the protected values.",
        indent=0.25,
    )
    add_body(
        doc,
        marker,
        "The positive application span proves this specific trace reached APM; the "
        "missing health span then proves filtering. If an organization hides a tag in "
        "the interface, use `agent-traces.out` as the authoritative attribute check.",
        indent=0.25,
    )

    add_numbered_step(doc, marker, step, "Confirm that the unchanged host-metrics pipeline still reaches the backend.")
    add_code(
        doc,
        marker,
        code_template,
        "jq -r '\n"
        "  .resourceMetrics[].resource.attributes[]\n"
        "  | select(.key == \"host.name\")\n"
        "  | .value.stringValue\n"
        "' agent-metrics.out | sort -u",
    )
    add_body(
        doc,
        marker,
        "Open **Infrastructure > Hosts**, find the detected host, and confirm that CPU, "
        "memory, load, or network data is current. Use the platform-specific "
        "Infrastructure navigator for EC2, Azure, GCP, or Kubernetes hosts.",
        indent=0.25,
    )

    add_lead(doc, marker, "Take-home: validate transformed logs locally")
    add_body(
        doc,
        marker,
        "No HEC token is required. With the updated Agent running, generate five JSON "
        "quote logs and compare `quotes.log` with `agent-logs.out`.",
    )
    add_code(doc, marker, code_template, "../loadgen -logs -json -count 5\nhead -n 1 quotes.log")
    add_bullet(doc, marker, bullet_template, "`level`, `message`, `movie`, and `timestamp` become log attributes.")
    add_bullet(doc, marker, bullet_template, "DEBUG, INFO, WARN, and ERROR map to severity numbers 5, 9, 13, and 17.")
    add_bullet(doc, marker, bullet_template, "Only `com.splunk.sourcetype`, `host.name`, and `otelcol.service.mode` remain as resource attributes.")

    add_lead(doc, marker, "Take-home: extend collection")
    add_link_line(
        doc,
        marker,
        "Splunk HEC exporter documentation",
        "https://help.splunk.com/en/splunk-observability-cloud/manage-data/splunk-distribution-of-the-opentelemetry-collector/get-started-with-the-splunk-distribution-of-the-opentelemetry-collector/collector-components/exporters/splunk-hec-exporter",
        lead="Logs to a non-production Splunk Platform instance: ",
    )
    add_link_line(
        doc,
        marker,
        "Introduction to Splunk Log Observer Connect",
        "https://help.splunk.com/splunk-observability-cloud/manage-data/view-splunk-platform-logs/introduction-to-splunk-log-observer-connect",
        lead="Correlate Platform logs with metrics and traces: ",
    )
    add_link_line(
        doc,
        marker,
        "Get data into Splunk APM AlwaysOn Profiling",
        "https://help.splunk.com/en/splunk-observability-cloud/monitor-application-performance/alwayson-profiling/get-data-into-splunk-apm-alwayson-profiling",
        lead="Profiling: ",
    )

    add_note(
        doc,
        marker,
        note_template,
        "A free Splunk Observability Cloud organization does not, by itself, provide a "
        "Splunk Platform HEC endpoint. Use a separate non-production Splunk Enterprise "
        "or Splunk Cloud Platform environment for the log-export take-home. Keep the "
        "final `agent_config.yaml` without secrets and review credentials, processor "
        "ordering, network exposure, and data volume before production use.",
    )


build_exercise_one(markers[0])
build_exercise_two(markers[1])
build_exercise_three(markers[2])

# Keep the cached TOC page numbers accurate in viewers that do not refresh fields.
# The PAGEREF fields themselves remain intact and Word refreshes them when opened.
toc_cached_pages = {
    11: 6,
    12: 6,
    13: 6,
    14: 8,
    15: 8,
    16: 8,
    17: 10,
    18: 10,
    19: 10,
}
for paragraph_index, page_number in toc_cached_pages.items():
    text_nodes = doc.paragraphs[paragraph_index]._p.xpath(".//w:t")
    cached_number_nodes = [
        node for node in text_nodes if (node.text or "").strip().isdigit()
    ]
    assert cached_number_nodes, doc.paragraphs[paragraph_index].text
    cached_number_nodes[-1].text = str(page_number)

# Ask Word to refresh PAGE/PAGEREF fields when opened.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.title = "OBS1184 - Advanced OpenTelemetry Collector Lab Guide"
doc.core_properties.subject = "Splunk .conf26 hands-on lab guide"
doc.core_properties.author = "Splunk"
doc.core_properties.comments = "Filled from the .conf26 lab guide template."

doc.save(OUTPUT)
print(OUTPUT)
